from typing import List, Literal, Union, Optional, Annotated
from pydantic import BaseModel

from tempfile import NamedTemporaryFile

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from openpyxl import Workbook
from openpyxl.utils import get_column_letter, coordinate_to_tuple
from openpyxl.worksheet.table import Table, TableStyleInfo

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from uuid import uuid4
from time import time
from langchain_core.tools import (
    tool,
)  # oppure `from langchain.tools import tool` in base alla versione


class TextData(BaseModel):
    """
    It's a block of simple text.
    Attributes:
        text: Text content to be used, for example for paragraphs,
              headings, or other textual components in a document.
    """

    text: str


class TableData(BaseModel):
    """
    It's a table structure for use in report or document generators (e.g., PDF, Word).
    Attributes:
        data: A matrix of strings representing the table content. Each sub-list is a row, each element a cell.
        col_widths: An optional list of column widths, expressed in units dependent on the rendering backend (e.g., points). If None, width is determined automatically.
        header: Indicates if the first row of `data` should be interpreted as the table header row.
    """

    data: List[List[str]]
    col_widths: Optional[List[float]] = None
    header: bool = True


class Component(BaseModel):
    """
    Single PDF/Word component.

    type:
      - "h1"    -> Heading1
      - "h2"    -> Heading2
      - "par"   -> Paragraph
      - "table" -> Table
    """

    type: Literal["h1", "h2", "par", "table"]
    data: Union[TextData, TableData]


class CreateSheetData(BaseModel):
    """
    Create a new sheet in the workbook.
    """

    name: str
    # Optional: position index (0-based). If None, append at the end.
    index: Optional[int] = None


class CellValueData(BaseModel):
    """
    Set a single cell value.
    You can use this both to write real data and to create a template placeholder.
    Example of placeholder: "{{customer_name}}"
    """

    sheet: str = "Sheet1"
    cell: str  # e.g. "B3"
    value: Union[str, float, int, bool, None]


class CellFormulaData(BaseModel):
    """
    Set a formula in a single cell.
    Formula must be a valid Excel formula WITHOUT leading '='.
    Example: "SUM(A1:A10)"
    """

    sheet: str = "Sheet1"
    cell: str
    formula: str


class TableData(BaseModel):
    """
    Insert a table starting from a top-left cell.
    """

    sheet: str = "Sheet1"
    start_cell: str  # e.g. "A1"
    data: List[List[Union[str, float, int, bool, None]]]
    with_header: bool = True
    auto_filter: bool = True
    table_style: Optional[str] = "TableStyleMedium9"


class XlsxOperation(BaseModel):
    """
    Single operation on the workbook.

    type:
      - "create_sheet" -> Create new sheet
      - "cell_value"   -> Write a single cell value
      - "cell_formula" -> Write a single cell formula
      - "table"        -> Insert a formatted table
    """

    type: Literal["create_sheet", "cell_value", "cell_formula", "table"]
    data: Union[CreateSheetData, CellValueData, CellFormulaData, TableData]


@tool()
def generate_docx(
    file_name: Annotated[str, "The name of file to generate"],
    components: Annotated[
        List[Component], "The list of component to use in order to create the Word file"
    ]
) -> str:
    """
    Generate a .docx Word document from a list of components.

    Components:
    - h1: main heading
    - h2: sub heading
    - par: paragraph
    - table: table with optional header

    Returns:
        Path of the generated .docx file in a temporary directory.
    """
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for comp in components:
        if comp.type == "h1":
            data: TextData = comp.data  # type: ignore
            p = doc.add_heading(level=1)
            run = p.add_run(data.text)
            run.font.size = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif comp.type == "h2":
            data: TextData = comp.data  # type: ignore
            p = doc.add_heading(level=2)
            run = p.add_run(data.text)
            run.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif comp.type == "par":
            data: TextData = comp.data  # type: ignore
            p = doc.add_paragraph()
            run = p.add_run(data.text)
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        elif comp.type == "table":
            data: TableData = comp.data  # type: ignore
            if not data.data:
                continue

            rows = len(data.data)
            cols = max(len(r) for r in data.data)

            table = doc.add_table(rows=rows, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.style = "Table Grid"

            for r_idx, row in enumerate(data.data):
                for c_idx, value in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = "" if value is None else str(value)

            # Header in bold if specified
            if data.header and len(data.data) > 0:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    # Temporary file
    tmp_file = NamedTemporaryFile(delete=False, suffix=".docx")
    docx_path = tmp_file.name
    tmp_file.close()

    doc.save(docx_path)

    return dict(path=docx_path, file_name=file_name, creation_time=time())


@tool()
def generate_xlsx(
    file_name: Annotated[str, "The name of file to generate"],
    operations: Annotated[
        List[XlsxOperation],
        "List of excel operation to perform in order to create an excel file",
    ],
) -> str:
    """
    Generate an .xlsx file based on a list of operations.

    Supported operations:
    - create_sheet: create a new sheet with a specific name
    - cell_value: set a cell value (text, number, boolean, or even placeholder like "{{name}}")
    - cell_formula: set a cell formula (without leading '=')
    - table: insert a table with optional header and style

    This tool can be used both to generate final reports and to create templates
    that other tools or users can later populate.

    Returns the path of the generated Excel file in a temporary directory,
    the name of file, and the creation timestamp in float.
    """
    wb = Workbook()
    # By default Workbook creates one sheet named "Sheet"
    default_sheet = wb.active
    default_sheet.title = "Sheet1"

    # Keep a counter for table names (Excel requires unique table names)
    table_counter = 1

    def get_sheet(name: str):
        if name in wb.sheetnames:
            return wb[name]
        # If sheet not exists, create it automatically
        return wb.create_sheet(title=name)

    for op in operations:
        if op.type == "create_sheet":
            data: CreateSheetData = op.data  # type: ignore
            # If already exists, skip
            if data.name in wb.sheetnames:
                continue
            if data.index is not None:
                wb.create_sheet(title=data.name, index=data.index)
            else:
                wb.create_sheet(title=data.name)

        elif op.type == "cell_value":
            data: CellValueData = op.data  # type: ignore
            ws = get_sheet(data.sheet)
            ws[data.cell] = data.value

        elif op.type == "cell_formula":
            data: CellFormulaData = op.data  # type: ignore
            ws = get_sheet(data.sheet)
            # Add leading '=' if not provided
            formula_str = data.formula
            if not formula_str.startswith("="):
                formula_str = "=" + formula_str
            ws[data.cell] = formula_str

        elif op.type == "table":
            data: TableData = op.data  # type: ignore
            ws = get_sheet(data.sheet)

            if not data.data:
                continue  # avoid empty tables to not explode

            # Write data to sheet
            start_row, start_col = coordinate_to_tuple(data.start_cell)
            max_cols = max(len(row) for row in data.data)
            max_rows = len(data.data)

            for r_idx, row in enumerate(data.data):
                for c_idx, value in enumerate(row):
                    ws.cell(
                        row=start_row + r_idx,
                        column=start_col + c_idx,
                        value=value,
                    )

            # Compute table range
            end_row = start_row + max_rows - 1
            end_col = start_col + max_cols - 1
            start_col_letter = get_column_letter(start_col)
            end_col_letter = get_column_letter(end_col)
            ref = f"{start_col_letter}{start_row}:{end_col_letter}{end_row}"

            # Create table with unique name
            table_name = f"Table_{table_counter}_{uuid4().hex[:6]}"
            table_counter += 1

            table = Table(displayName=table_name, ref=ref)

            if data.table_style is not None:
                style = TableStyleInfo(
                    name=data.table_style,
                    showFirstColumn=False,
                    showLastColumn=False,
                    showRowStripes=True,
                    showColumnStripes=False,
                )
                table.tableStyleInfo = style

            # AutoFilter handled by the table itself
            ws.add_table(table)

    # Create temporary file
    tmp_file = NamedTemporaryFile(delete=False, suffix=".xlsx")
    xlsx_path = tmp_file.name
    tmp_file.close()

    wb.save(xlsx_path)

    return dict(path=xlsx_path, file_name=file_name, creation_time=time())


@tool()
def generate_pdf(
    file_name: Annotated[str, "The name of file to generate"],
    components: Annotated[
        List[Component], "The list of component to use in order to create the pdf file"
    ],
) -> str:
    """
    Generate an A4 PDF from a list of components.

    Each component can be:
    - h1: main heading
    - h2: sub heading
    - par: paragraph
    - table: table with data

    Returns the path of the generated PDF file in a temporary directory,
    the name of file, and the creation timestamp in float.
    """
    # Temporary file
    tmp_file = NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf_path = tmp_file.name
    tmp_file.close()

    # Document
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        leftMargin=40,
        rightMargin=40,
        topMargin=50,
        bottomMargin=50,
    )

    # Styles
    styles = getSampleStyleSheet()

    # Override / add custom styles
    styles.add(
        ParagraphStyle(
            name="CustomHeading1",
            parent=styles["Heading1"],
            fontSize=20,
            leading=24,
            spaceAfter=12,
            textColor=colors.HexColor("#222222"),
        )
    )

    styles.add(
        ParagraphStyle(
            name="CustomHeading2",
            parent=styles["Heading2"],
            fontSize=16,
            leading=20,
            spaceBefore=6,
            spaceAfter=8,
            textColor=colors.HexColor("#333333"),
        )
    )

    styles.add(
        ParagraphStyle(
            name="CustomParagraph",
            parent=styles["BodyText"],
            fontSize=10.5,
            leading=14,
            spaceAfter=6,
        )
    )

    story = []

    for comp in components:
        if comp.type == "h1":
            data: TextData = comp.data  # type: ignore
            story.append(Paragraph(data.text, styles["CustomHeading1"]))
            story.append(Spacer(1, 8))

        elif comp.type == "h2":
            data: TextData = comp.data  # type: ignore
            story.append(Paragraph(data.text, styles["CustomHeading2"]))
            story.append(Spacer(1, 6))

        elif comp.type == "par":
            data: TextData = comp.data  # type: ignore
            story.append(Paragraph(data.text, styles["CustomParagraph"]))
            story.append(Spacer(1, 4))

        elif comp.type == "table":
            data: TableData = comp.data  # type: ignore
            tbl = Table(data.data, colWidths=data.col_widths)

            # Basic table styling
            table_style = [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]

            if data.header and len(data.data) > 0:
                table_style.extend(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#000000")),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ]
                )

            tbl.setStyle(TableStyle(table_style))

            story.append(tbl)
            story.append(Spacer(1, 8))

    # Build PDF
    doc.build(story)

    # Return file info
    return dict(path=pdf_path, file_name=file_name, creation_time=time())

def load_tools():
    """
    Load tools.
    """
    return [generate_pdf, generate_xlsx, generate_docx]
